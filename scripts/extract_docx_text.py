from pathlib import Path
import sys

from docx import Document


def main() -> None:
    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    document = Document(source)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style else ""
            lines.append(f"[{style}] {text}" if style else text)

    for table_index, table in enumerate(document.tables, start=1):
        lines.append(f"\n[TABLE {table_index}]")
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            lines.append(" | ".join(cells))

    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text("\n".join(lines), encoding="utf-8")
    print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} chars={sum(len(line) for line in lines)}")


if __name__ == "__main__":
    main()
